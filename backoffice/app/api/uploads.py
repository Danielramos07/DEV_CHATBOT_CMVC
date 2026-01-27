from flask import Blueprint, request, jsonify
from flask import send_file
from ..db import get_conn
from ..services.retreival import build_faiss_index
from ..services.rag import index_pdf_documents
from ..services.text import normalizar_idioma
from ..config import Config
from ..services.video_service import can_start_new_video_job
import os
import traceback
import re
import unicodedata
import io
import zipfile
import xml.etree.ElementTree as ET

# Importações opcionais para evitar erros se os pacotes não estiverem instalados
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

app = Blueprint('uploads', __name__)

_FAQ_KEY_MAP = {
    "designacao da faq": "designacao",
    "designacao faq": "designacao",
    "designacao": "designacao",
    "titulo da faq": "designacao",
    "titulo faq": "designacao",
    "titulo": "designacao",
    "pergunta": "pergunta",
    "questao": "pergunta",
    "resposta": "resposta",
    "categoria": "categoria",
    "idioma": "idioma",
    "identificador": "identificador",
    "id": "identificador",
    "documentos associados": "links_documentos",
    "links de documentos": "links_documentos",
    "links documentos": "links_documentos",
    "documentos": "links_documentos",
    "links": "links_documentos",
    "a quem se destina": "serve_text",
    "para que serve": "serve_text",
    "serve": "serve_text",
    "identificador codigo da faq": "identificador",
    "codigo da faq": "identificador",
    "identificador da faq": "identificador",
    "o que tem que fazer e quais os documentos necessarios": "resposta",
    "a quem se destina e para que serve este procedimento": "serve_text",
}

_REQUIRED_FAQ_FIELDS = ("designacao", "pergunta", "resposta")


def _normalize_faq_key(raw: str) -> str:
    value = (raw or "").strip().lower().replace("\u2019", "'").replace("\u2018", "'")
    value = unicodedata.normalize("NFKD", value)
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    value = value.replace(":", " ")
    value = value.replace("/", " ")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _extract_docx_pairs(doc):
    pairs = []
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text
                value = row.cells[1].text
                pairs.append((key, value))
    for para in doc.paragraphs:
        text = para.text.strip()
        if ":" in text:
            key, value = text.split(":", 1)
            if key.strip() and value.strip():
                pairs.append((key, value))
    return pairs


def _parse_faq_pairs(pairs):
    data = {}
    last_canonical = None
    for key, value in pairs:
        raw_value = (value or "").strip()
        normalized_key = _normalize_faq_key(key or "")
        if not normalized_key:
            if last_canonical and raw_value:
                existing = data.get(last_canonical, "")
                data[last_canonical] = (existing + "\n" + raw_value).strip()
            continue
        canonical = _FAQ_KEY_MAP.get(normalized_key)
        if not canonical:
            if normalized_key.startswith("designacao"):
                canonical = "designacao"
            elif normalized_key.startswith("pergunta"):
                canonical = "pergunta"
            elif normalized_key.startswith("questao"):
                canonical = "pergunta"
            elif normalized_key.startswith("resposta"):
                canonical = "resposta"
            elif normalized_key.startswith("identificador"):
                canonical = "identificador"
        if not canonical or not raw_value:
            continue
        last_canonical = canonical
        existing = data.get(canonical)
        data[canonical] = raw_value if not existing else (existing + "\n" + raw_value).strip()
    return data


def _parse_docx_faq_data(file_obj):
    doc = docx.Document(file_obj)
    return _parse_faq_pairs(_extract_docx_pairs(doc))


def _extract_odt_pairs(xml_bytes: bytes):
    ns = {
        "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
        "table": "urn:oasis:names:tc:opendocument:xmlns:table:1.0",
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
    }
    root = ET.fromstring(xml_bytes)
    pairs = []
    for table in root.findall(".//table:table", ns):
        for row in table.findall("table:table-row", ns):
            cells = []
            for cell in row.findall("table:table-cell", ns):
                parts = []
                for p in cell.findall(".//text:p", ns):
                    chunk = "".join(p.itertext()).strip()
                    if chunk:
                        parts.append(chunk)
                cells.append("\n".join(parts).strip())
            if len(cells) >= 2 and any(cells[:2]):
                pairs.append((cells[0], cells[1]))
    # Fallback to paragraphs with "key: value" style.
    for para in root.findall(".//text:p", ns):
        text = "".join(para.itertext()).strip()
        if ":" in text:
            key, value = text.split(":", 1)
            if key.strip() and value.strip():
                pairs.append((key, value))
    return pairs


def _parse_odt_faq_data(file_bytes: bytes):
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
        if "content.xml" not in zf.namelist():
            return {}
        xml_bytes = zf.read("content.xml")
    return _parse_faq_pairs(_extract_odt_pairs(xml_bytes))


def _parse_faq_upload(file_storage):
    file_bytes = file_storage.read()
    if not file_bytes:
        return {}
    filename = (file_storage.filename or "").lower()
    if filename.endswith(".odt"):
        return _parse_odt_faq_data(file_bytes)
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            names = set(zf.namelist())
        if "content.xml" in names:
            return _parse_odt_faq_data(file_bytes)
    except zipfile.BadZipFile:
        pass
    return _parse_docx_faq_data(io.BytesIO(file_bytes))


@app.route("/upload-pdf", methods=["POST"])
def upload_pdf():
    if not PDF_AVAILABLE:
        return jsonify({"success": False, "error": "PyPDF2 não está instalado."}), 500
    
    conn = get_conn()
    cur = conn.cursor()
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "Ficheiro não enviado."}), 400
    files = request.files.getlist('file')
    chatbot_id = request.form.get("chatbot_id")
    if not chatbot_id:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400
    uploaded_pdf_ids = []
    try:
        if not os.path.exists(Config.PDF_STORAGE_PATH):
            os.makedirs(Config.PDF_STORAGE_PATH, exist_ok=True)
        for file in files:
            filename = file.filename
            if not filename.lower().endswith('.pdf'):
                return jsonify({"success": False, "error": "Apenas ficheiros PDF são permitidos."}), 400
            # Store per chatbot to avoid name collisions
            chatbot_dir = os.path.join(Config.PDF_STORAGE_PATH, f"chatbot_{chatbot_id}")
            os.makedirs(chatbot_dir, exist_ok=True)
            file_path = os.path.join(chatbot_dir, filename)
            file.save(file_path)
            file.close()
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if reader.is_encrypted:
                    return jsonify({"success": False, "error": f"O PDF '{filename}' está protegido por senha."}), 400
                if not reader.pages:
                    return jsonify({"success": False, "error": f"O PDF '{filename}' está vazio ou corrompido."}), 400
                cur.execute(
                    "INSERT INTO pdf_documents (chatbot_id, filename, file_path) VALUES (%s, %s, %s) RETURNING pdf_id",
                (chatbot_id, filename, file_path)
            )
            pdf_id = cur.fetchone()[0]
            uploaded_pdf_ids.append(pdf_id)
        conn.commit()
        rag_indexed = False
        try:
            index_pdf_documents(chatbot_id=int(chatbot_id), pdf_ids=uploaded_pdf_ids)
            rag_indexed = True
        except Exception:
            traceback.print_exc()
        return jsonify({
            "success": True,
            "pdf_ids": uploaded_pdf_ids,
            "rag_indexed": rag_indexed,
            "message": "PDF(s) carregado(s) com sucesso."
        })
    except PyPDF2.errors.PdfReadError:
        return jsonify({"success": False, "error": "Erro ao ler o PDF. Verifique se o arquivo não está corrompido."}), 400
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        cur.close()
        conn.close()


@app.route("/pdfs/<int:pdf_id>", methods=["GET"])
def get_pdf(pdf_id: int):
    """Serve an uploaded PDF from pdf_documents by id."""
    conn = get_conn()
    cur = conn.cursor()
    try:
        cur.execute("SELECT file_path FROM pdf_documents WHERE pdf_id = %s", (pdf_id,))
        row = cur.fetchone()
        if not row or not row[0]:
            return jsonify({"success": False, "error": "PDF não encontrado."}), 404
        file_path = row[0]
        if not os.path.isfile(file_path):
            return jsonify({"success": False, "error": "Ficheiro PDF não existe no servidor."}), 404
        return send_file(file_path, mimetype="application/pdf", as_attachment=False)
    finally:
        cur.close()
        conn.close()

@app.route("/rebuild-rag", methods=["POST"])
def rebuild_rag():
    data = request.get_json(silent=True) or {}
    chatbot_id = data.get("chatbot_id")
    try:
        if chatbot_id:
            chunks = index_pdf_documents(chatbot_id=int(chatbot_id))
        else:
            chunks = index_pdf_documents()
        return jsonify({"success": True, "chunks": chunks})
    except Exception as exc:
        traceback.print_exc()
        return jsonify({"success": False, "error": str(exc)}), 500

@app.route("/upload-faq-docx", methods=["POST"])
def upload_faq_docx():
    if not DOCX_AVAILABLE:
        return jsonify({"success": False, "error": "python-docx não está instalado."}), 500
    
    conn = get_conn()
    cur = conn.cursor()
    # Accept both "file" and "files" (some forms send "files" even for a single docx)
    if 'file' in request.files:
        file = request.files['file']
    elif 'files' in request.files:
        files = request.files.getlist('files')
        file = files[0] if files else None
    else:
        return jsonify({"success": False, "error": "Ficheiro não enviado."}), 400
    if not file:
        return jsonify({"success": False, "error": "Ficheiro não enviado."}), 400
    chatbot_id_raw = request.form.get("chatbot_id")
    if not chatbot_id_raw:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400
    # If a video job is active, only block when at least one target chatbot has video enabled
    if not can_start_new_video_job():
        try:
            if chatbot_id_raw == "todos":
                cur.execute("SELECT 1 FROM chatbot WHERE video_enabled = TRUE LIMIT 1")
                any_video = cur.fetchone() is not None
            else:
                cur.execute("SELECT video_enabled FROM chatbot WHERE chatbot_id = %s", (int(chatbot_id_raw),))
                r = cur.fetchone()
                any_video = bool(r[0]) if r else False
            if any_video:
                return jsonify({"success": False, "busy": True, "error": "Já existe um vídeo a ser gerado. Aguarde que termine."}), 409
        except Exception:
            return jsonify({"success": False, "busy": True, "error": "Já existe um vídeo a ser gerado. Aguarde que termine."}), 409
    try:
        dados = _parse_faq_upload(file)
        missing = [f for f in _REQUIRED_FAQ_FIELDS if not dados.get(f)]
        if missing:
            raise Exception("Faltam campos obrigatórios: designação, questão ou resposta.")
        designacao = dados.get("designacao")
        identificador = dados.get("identificador") or ""
        pergunta = dados.get("pergunta")
        resposta = dados.get("resposta")
        categoria = dados.get("categoria")
        idioma_lido = dados.get("idioma", "Português")
        idioma = normalizar_idioma(idioma_lido)
        links_documentos = dados.get("links_documentos", "")
        serve_text = dados.get("serve_text")
        chatbot_ids = []
        if chatbot_id_raw == "todos":
            cur.execute("SELECT chatbot_id FROM chatbot")
            chatbot_ids = [row[0] for row in cur.fetchall()]
        else:
            chatbot_ids = [int(chatbot_id_raw)]
        for chatbot_id in chatbot_ids:
            cur.execute("""
                SELECT faq_id FROM faq
                WHERE chatbot_id = %s AND designacao = %s AND pergunta = %s AND resposta = %s AND idioma = %s
            """, (chatbot_id, designacao, pergunta, resposta, idioma))
            if cur.fetchone():
                continue
            cur.execute("""
                INSERT INTO faq (chatbot_id, designacao, identificador, pergunta, resposta, idioma, links_documentos, serve_text)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING faq_id
            """, (
                chatbot_id,
                designacao,
                (identificador or "").strip() or None,
                pergunta,
                resposta,
                idioma,
                links_documentos,
                (serve_text or "").strip() or None,
            ))
            faq_id = cur.fetchone()[0]
            if categoria:
                cur.execute("SELECT categoria_id FROM categoria WHERE nome ILIKE %s", (categoria,))
                result = cur.fetchone()
                if result:
                    cur.execute("UPDATE faq SET categoria_id = %s WHERE faq_id = %s", (result[0], faq_id))
            if links_documentos:
                for link in re.split(r"[,\n]+", links_documentos):
                    link = link.strip()
                    if link:
                        cur.execute(
                            "INSERT INTO faq_documento (faq_id, link) VALUES (%s, %s)",
                            (faq_id, link)
                        )
        conn.commit()
        build_faiss_index()
        return jsonify({"success": True, "message": "FAQ e links inseridos com sucesso."})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        cur.close()
        conn.close()

@app.route("/upload-faq-docx-multiplos", methods=["POST"])
def upload_faq_docx_multiplos():
    if not DOCX_AVAILABLE:
        return jsonify({"success": False, "error": "python-docx não está instalado."}), 500
    
    conn = get_conn()
    cur = conn.cursor()
    if 'files' not in request.files and 'file' not in request.files:
        return jsonify({"success": False, "error": "Ficheiros não enviados."}), 400
    chatbot_id_raw = request.form.get("chatbot_id")
    if not chatbot_id_raw:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400
    # If a video job is active, only block when at least one target chatbot has video enabled
    if not can_start_new_video_job():
        try:
            cur2 = conn.cursor()
            if chatbot_id_raw == "todos":
                cur2.execute("SELECT 1 FROM chatbot WHERE video_enabled = TRUE LIMIT 1")
                any_video = cur2.fetchone() is not None
            else:
                cur2.execute("SELECT video_enabled FROM chatbot WHERE chatbot_id = %s", (int(chatbot_id_raw),))
                r = cur2.fetchone()
                any_video = bool(r[0]) if r else False
            cur2.close()
            if any_video:
                return jsonify({"success": False, "busy": True, "error": "Já existe um vídeo a ser gerado. Aguarde que termine."}), 409
        except Exception:
            return jsonify({"success": False, "busy": True, "error": "Já existe um vídeo a ser gerado. Aguarde que termine."}), 409
    files = request.files.getlist('files') or request.files.getlist('file')
    total_inseridas = 0
    erros = []
    for file in files:
        try:
            dados = _parse_faq_upload(file)
            missing = [f for f in _REQUIRED_FAQ_FIELDS if not dados.get(f)]
            if missing:
                raise Exception("Faltam campos obrigatórios: designação, questão ou resposta.")
            designacao = dados.get("designacao")
            identificador = dados.get("identificador") or ""
            pergunta = dados.get("pergunta")
            resposta = dados.get("resposta")
            categoria = dados.get("categoria")
            idioma_lido = dados.get("idioma", "Português")
            idioma = normalizar_idioma(idioma_lido)
            links_documentos = dados.get("links_documentos", "")
            serve_text = dados.get("serve_text")
            chatbot_ids = []
            if chatbot_id_raw == "todos":
                cur.execute("SELECT chatbot_id FROM chatbot")
                chatbot_ids = [row[0] for row in cur.fetchall()]
            else:
                chatbot_ids = [int(chatbot_id_raw)]
            for chatbot_id in chatbot_ids:
                cur.execute("""
                    SELECT faq_id FROM faq
                    WHERE chatbot_id = %s AND designacao = %s AND pergunta = %s AND resposta = %s AND idioma = %s
                """, (chatbot_id, designacao, pergunta, resposta, idioma))
                if cur.fetchone():
                    continue
                cur.execute("""
                    INSERT INTO faq (chatbot_id, designacao, identificador, pergunta, resposta, idioma, links_documentos, serve_text)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING faq_id
                """, (
                    chatbot_id,
                    designacao,
                    (identificador or "").strip() or None,
                    pergunta,
                    resposta,
                    idioma,
                    links_documentos,
                    (serve_text or "").strip() or None,
                ))
                faq_id = cur.fetchone()[0]
                if categoria:
                    cur.execute("SELECT categoria_id FROM categoria WHERE nome ILIKE %s", (categoria,))
                    result = cur.fetchone()
                    if result:
                        cur.execute("UPDATE faq SET categoria_id = %s WHERE faq_id = %s", (result[0], faq_id))
                if links_documentos:
                    for link in re.split(r"[,\n]+", links_documentos):
                        link = link.strip()
                        if link:
                            cur.execute(
                                "INSERT INTO faq_documento (faq_id, link) VALUES (%s, %s)",
                                (faq_id, link)
                            )
                total_inseridas += 1
        except Exception as e:
            erros.append(str(e))
            conn.rollback()
    conn.commit()
    build_faiss_index()
    return jsonify({"success": True, "inseridas": total_inseridas, "erros": erros})

